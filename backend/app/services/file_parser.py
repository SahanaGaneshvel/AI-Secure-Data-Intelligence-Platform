"""
File parsing service - extract text from various file formats
Supports: .txt, .log, .pdf, .docx
"""
import io
from typing import Optional, Tuple
from enum import Enum

class FileType(str, Enum):
    TXT = "txt"
    LOG = "log"
    PDF = "pdf"
    DOCX = "docx"
    UNKNOWN = "unknown"

class FileParserService:
    """Extract text content from uploaded files"""

    @staticmethod
    def detect_file_type(filename: str) -> FileType:
        """Detect file type from filename"""
        extension = filename.lower().split('.')[-1] if '.' in filename else ''

        if extension in ['txt']:
            return FileType.TXT
        elif extension in ['log']:
            return FileType.LOG
        elif extension in ['pdf']:
            return FileType.PDF
        elif extension in ['docx', 'doc']:
            return FileType.DOCX
        else:
            return FileType.UNKNOWN

    @staticmethod
    def parse_txt(content: bytes) -> str:
        """Parse .txt file"""
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1 if utf-8 fails
            return content.decode('latin-1', errors='ignore')

    @staticmethod
    def parse_log(content: bytes) -> str:
        """Parse .log file (same as txt)"""
        return FileParserService.parse_txt(content)

    @staticmethod
    def parse_pdf(content: bytes) -> str:
        """Parse .pdf file using pdfplumber (fallback to PyPDF2)"""
        text_parts = []

        try:
            # Try pdfplumber first (better text extraction)
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            if text_parts:
                return '\n'.join(text_parts)
        except Exception as e:
            print(f"pdfplumber failed: {e}, trying PyPDF2...")

        try:
            # Fallback to PyPDF2
            import PyPDF2

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            if text_parts:
                return '\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")

        if not text_parts:
            raise ValueError("PDF parsing failed: no text extracted")

        return '\n'.join(text_parts)

    @staticmethod
    def parse_docx(content: bytes) -> str:
        """Parse .docx file using python-docx"""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            if not text_parts:
                raise ValueError("DOCX parsing failed: no text extracted")

            return '\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")

    @staticmethod
    def parse_file(filename: str, content: bytes) -> Tuple[str, FileType]:
        """
        Parse file and extract text content

        Returns:
            (extracted_text, file_type)
        """
        file_type = FileParserService.detect_file_type(filename)

        if file_type == FileType.TXT:
            return FileParserService.parse_txt(content), file_type
        elif file_type == FileType.LOG:
            return FileParserService.parse_log(content), file_type
        elif file_type == FileType.PDF:
            return FileParserService.parse_pdf(content), file_type
        elif file_type == FileType.DOCX:
            return FileParserService.parse_docx(content), file_type
        else:
            raise ValueError(f"Unsupported file type: {filename}")

    @staticmethod
    def validate_file_size(content: bytes, max_size_mb: int = 10) -> None:
        """Validate file size"""
        size_mb = len(content) / (1024 * 1024)
        if size_mb > max_size_mb:
            raise ValueError(f"File too large: {size_mb:.2f}MB (max: {max_size_mb}MB)")
